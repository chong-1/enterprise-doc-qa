"""Phase 3 单元测试：文档解析 / 文本清洗 / 分块 / 文件存储。"""

from io import BytesIO

import pytest

from app.services.document.chunker import split_text_into_chunks
from app.services.document.loader import delete_file, read_file, save_file
from app.services.document.parser import clean_text, parse_document


# ========== 文本清洗 ==========

class TestCleanText:
    def test_压缩空行(self):
        text = "第一行\n\n\n\n第二行\n\n第三行\n"
        # 连续空行压缩为单个空行（保留段间分隔）
        assert clean_text(text) == "第一行\n\n第二行\n\n第三行"

    def test_去除零宽字符(self):
        text = "你好​世界﻿测试"
        assert clean_text(text) == "你好世界测试"

    def test_每行去首尾空白(self):
        text = "  你好  \n  世界  "
        assert clean_text(text) == "你好\n世界"

    def test_空文本(self):
        assert clean_text("") == ""
        assert clean_text("  \n\n ") == ""


# ========== 解析器 ==========

class TestParser:
    def test_txt(self):
        assert parse_document("你好世界\n第二行".encode("utf-8"), "txt") == "你好世界\n第二行"

    def test_txt_gbk回退(self):
        content = "中文GBK编码".encode("gbk")
        assert parse_document(content, "txt") == "中文GBK编码"

    def test_markdown提取文本(self):
        md = "# 一级标题\n\n**加粗** 普通文本 [链接](https://example.com)\n\n```python\nprint('hi')\n```\n"
        result = parse_document(md.encode("utf-8"), "md")
        assert "一级标题" in result
        assert "加粗" in result  # 加粗标记被剥离
        assert "https://example.com" not in result  # 链接 URL 不进入文本
        assert "print('hi')" in result  # 代码块保留

    def test_docx段落和表格(self):
        import docx

        document = docx.Document()
        document.add_paragraph("第一段内容")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        buf = BytesIO()
        document.save(buf)
        result = parse_document(buf.getvalue(), "docx")
        assert "第一段内容" in result
        assert "A1\tB1" in result
        assert "A2\tB2" in result

    def test_xlsx保留表格结构(self):
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "销售数据"
        ws.append(["月份", "销售额"])
        ws.append(["一月", 100])
        buf = BytesIO()
        wb.save(buf)
        result = parse_document(buf.getvalue(), "xlsx")
        assert "销售数据" in result
        assert "月份\t销售额" in result
        assert "一月\t100" in result

    def test_pdf逐页提取(self):
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello PDF Page 1")
        page2 = doc.new_page()
        page2.insert_text((72, 72), "Hello PDF Page 2")
        result = parse_document(doc.tobytes(), "pdf")
        assert "Hello PDF Page 1" in result
        assert "Hello PDF Page 2" in result

    def test_不支持类型(self):
        from app.core.exceptions import BadRequestError

        with pytest.raises(BadRequestError):
            parse_document(b"x", "exe")


# ========== 分块器 ==========

class TestChunker:
    def test_短文本单块(self):
        chunks = split_text_into_chunks("这是很短的一句话。只有几个字。")
        assert len(chunks) == 1
        assert chunks[0].index == 0
        assert chunks[0].char_start == 0

    def test_长文本多块且不超上限(self):
        # 40 句，每句约 20 字 → 总 token 远超 512
        text = "".join(f"这是第{i}句测试文本，用于验证分块器能把长文档切成合理的块。\n" for i in range(40))
        chunks = split_text_into_chunks(text, chunk_size=100, chunk_overlap=20)
        assert len(chunks) >= 3
        # 每个 chunk 的 token 数 <= 100 + 单句余量（单句本身可能略超）
        from app.services.document.chunker import get_tokenizer

        tok = get_tokenizer()
        for c in chunks:
            n = len(tok.encode(c.text, add_special_tokens=False))
            assert n <= 100 + 30, f"chunk {c.index} 超出上限: {n} tokens"

    def test_overlap保留上下文(self):
        text = "".join(f"这是第{i}句测试文本。\n" for i in range(30))
        chunks = split_text_into_chunks(text, chunk_size=80, chunk_overlap=30)
        assert len(chunks) >= 2
        # 相邻 chunk 应共享重叠文本（前一个 chunk 尾部句子出现在后一个 chunk 开头）
        overlap = chunks[0].text.split()[-1]
        assert overlap in chunks[1].text

    def test_超长单句硬切(self):
        text = "这是一个非常长的句子，" * 50 + "结尾。"
        chunks = split_text_into_chunks(text, chunk_size=64, chunk_overlap=16)
        assert len(chunks) >= 2
        assert all(c.text for c in chunks)

    def test_chunk偏移与原文一致(self):
        text = "第一句话。第二句话。第三句话。第四句话。第五句话。" * 3
        chunks = split_text_into_chunks(text, chunk_size=50, chunk_overlap=10)
        for c in chunks:
            assert text[c.char_start : c.char_end].strip() == c.text
        # 覆盖全部内容（含重叠）
        covered = set()
        for c in chunks:
            covered.update(range(c.char_start, c.char_end))
        assert len(covered) == len(text) or len(covered) > len(text) * 0.9


# ========== 文件存储 ==========

class TestLoader:
    def test_roundtrip(self):
        rel = save_file(b"file-content-123", "测试.txt")
        assert read_file(rel) == b"file-content-123"
        delete_file(rel)
        from app.core.exceptions import NotFoundError

        with pytest.raises(NotFoundError):
            read_file(rel)

    def test_路径越界防护(self):
        from app.core.exceptions import NotFoundError

        with pytest.raises((ValueError, NotFoundError)):
            read_file("../outside.txt")

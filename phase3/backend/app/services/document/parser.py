"""多格式文档解析器 + 文本清洗。

支持格式：
- PDF    → PyMuPDF (fitz)，逐页提取文本
- Word   → python-docx，段落 + 表格
- Excel  → openpyxl，按 sheet/行保留表格结构
- Markdown → markdown-it-py，提取文本并保留标题层级
- TXT    → 直接读取（utf-8 / gbk 自动回退）
"""

import re
from io import BytesIO

from app.core.exceptions import BadRequestError

# 支持的文件扩展名（与文件类型校验共用）
SUPPORTED_FILE_TYPES = {"pdf", "docx", "xlsx", "md", "txt"}

# ========== 文本清洗 ==========

# 零宽字符 / 不可见控制字符（保留 \t \n \r）
_INVISIBLE_CHARS = re.compile(
    "[\u200b-\u200f\u202a-\u202e\u2060\u2066-\u2069\ufeff\x00-\x08\x0b\x0c\x0e-\x1f]"
)


def clean_text(text: str) -> str:
    """清洗提取出的原始文本：去乱码字符、压缩空行、规范化空白。"""
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _INVISIBLE_CHARS.sub("", text)
    # 每行去首尾空白，但保留表格内的 \t
    lines = [line.strip() for line in text.split("\n")]
    # 连续空行压缩为单个空行
    cleaned: list[str] = []
    blank = False
    for line in lines:
        if line:
            cleaned.append(line)
            blank = False
        elif not blank:
            cleaned.append("")
            blank = True
    return "\n".join(cleaned).strip()


# ========== 各格式解析器 ==========

def _parse_pdf(content: bytes) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        pages = []
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    finally:
        doc.close()


def _parse_docx(content: bytes) -> str:
    import docx

    document = docx.Document(BytesIO(content))
    parts: list[str] = []
    # 简化实现：先段落，后表格（表格内容已含行列信息，顺序损失可接受）
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _parse_xlsx(content: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(BytesIO(content), read_only=True, data_only=True)
    try:
        parts: list[str] = []
        for sheet in wb.worksheets:
            parts.append(f"# 工作表：{sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    finally:
        wb.close()


def _parse_markdown(content: bytes) -> str:
    from markdown_it import MarkdownIt

    md = MarkdownIt("commonmark")
    text = content.decode("utf-8", errors="replace")
    out: list[str] = []
    for token in md.parse(text):
        if token.type == "heading_open":
            # 保留标题层级信息（# 号），利于分块时识别章节边界
            out.append("#" * int(token.tag[1]) + " ")
        elif token.type == "inline":
            # inline token 的 children 中 text 类型才是纯文本（剔除链接/加粗标记）
            out.append("".join(c.content for c in token.children if c.type == "text"))
        elif token.type in ("fence", "code_block"):
            out.append(token.content)
    return "\n".join(out)


def _parse_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        # 中文文档常为 GBK 编码，utf-8 失败后回退
        return content.decode("gbk", errors="replace")


# 解析器注册表：扩展新格式时在此登记
_PARSERS: dict[str, callable] = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "xlsx": _parse_xlsx,
    "md": _parse_markdown,
    "txt": _parse_txt,
}


def parse_document(content: bytes, file_type: str) -> str:
    """按文件类型提取纯文本。不存在的类型抛 400。"""
    parser = _PARSERS.get(file_type)
    if parser is None:
        raise BadRequestError(f"不支持的文件类型: {file_type}")
    return parser(content)
